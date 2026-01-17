"""Property-based tests for resume parsing"""

import pytest
import tempfile
import os
from pathlib import Path
from hypothesis import given, strategies as st, settings, assume
from hypothesis import HealthCheck
import PyPDF2
from docx import Document

from ml.parsing.text_extractor import TextExtractor
from ml.parsing.section_identifier import SectionIdentifier
from ml.parsing.entity_extractor import EntityExtractor
from ml.parsing.resume_parser import ResumeParser
from backend.app.core.exceptions import ValidationException


# Test data generators
@st.composite
def resume_text(draw):
    """Generate realistic resume text"""
    sections = []
    
    # Header
    name = draw(st.text(min_size=5, max_size=30, alphabet=st.characters(whitelist_categories=('Lu', 'Ll', ' '))))
    sections.append(name)
    
    # Experience section
    if draw(st.booleans()):
        sections.append("\nWORK EXPERIENCE")
        num_jobs = draw(st.integers(min_value=1, max_value=3))
        for _ in range(num_jobs):
            title = draw(st.text(min_size=5, max_size=30))
            company = draw(st.text(min_size=5, max_size=30))
            sections.append(f"\n{title}\n{company}\nJan 2020 - Present")
    
    # Education section
    if draw(st.booleans()):
        sections.append("\nEDUCATION")
        degree = draw(st.sampled_from(['Bachelor', 'Master', 'PhD']))
        university = draw(st.text(min_size=10, max_size=40))
        sections.append(f"\n{degree} of Science\n{university}\n2016 - 2020")
    
    # Skills section
    if draw(st.booleans()):
        sections.append("\nSKILLS")
        num_skills = draw(st.integers(min_value=3, max_value=10))
        skills = [draw(st.text(min_size=3, max_size=15)) for _ in range(num_skills)]
        sections.append("\n" + ", ".join(skills))
    
    return "\n".join(sections)


@st.composite
def pdf_file(draw):
    """Generate a temporary PDF file with resume content"""
    content = draw(resume_text())
    assume(len(content) > 50)  # Ensure meaningful content
    
    # Create temporary PDF
    temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.pdf', delete=False)
    temp_file.close()
    
    # Write PDF content
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    
    c = canvas.Canvas(temp_file.name, pagesize=letter)
    
    # Split content into lines and write to PDF
    y_position = 750
    for line in content.split('\n'):
        if y_position < 50:
            c.showPage()
            y_position = 750
        c.drawString(50, y_position, line[:80])  # Limit line length
        y_position -= 15
    
    c.save()
    
    return temp_file.name, content


@st.composite
def docx_file(draw):
    """Generate a temporary DOCX file with resume content"""
    content = draw(resume_text())
    assume(len(content) > 50)  # Ensure meaningful content
    
    # Create temporary DOCX
    temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False)
    temp_file.close()
    
    # Write DOCX content
    doc = Document()
    for line in content.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    doc.save(temp_file.name)
    
    return temp_file.name, content


# Property 1: Resume format handling
@pytest.mark.property
@settings(max_examples=50, suppress_health_check=[HealthCheck.function_scoped_fixture])
@given(file_data=pdf_file())
def test_property_1_pdf_format_handling(file_data):
    """
    **Feature: talentflow-ai, Property 1: Resume format handling**
    **Validates: Requirements 1.1, 1.2**
    
    Property: System must successfully extract text from valid PDF files
    """
    file_path, expected_content = file_data
    
    try:
        extractor = TextExtractor()
        extracted_text = extractor.extract_text(file_path)
        
        # Verify text was extracted
        assert extracted_text is not None
        assert len(extracted_text) > 0
        assert isinstance(extracted_text, str)
    
    finally:
        # Cleanup
        if os.path.exists(file_path):
            os.unlink(file_path)


@pytest.mark.property
@settings(max_examples=50, suppress_health_check=[HealthCheck.function_scoped_fixture])
@given(file_data=docx_file())
def test_property_1_docx_format_handling(file_data):
    """
    **Feature: talentflow-ai, Property 1: Resume format handling**
    **Validates: Requirements 1.1, 1.2**
    
    Property: System must successfully extract text from valid DOCX files
    """
    file_path, expected_content = file_data
    
    try:
        extractor = TextExtractor()
        extracted_text = extractor.extract_text(file_path)
        
        # Verify text was extracted
        assert extracted_text is not None
        assert len(extracted_text) > 0
        assert isinstance(extracted_text, str)
    
    finally:
        # Cleanup
        if os.path.exists(file_path):
            os.unlink(file_path)


# Property 4: Minimum field extraction
@pytest.mark.property
@settings(max_examples=100)
@given(text=resume_text())
def test_property_4_minimum_field_extraction(text):
    """
    **Feature: talentflow-ai, Property 4: Minimum field extraction**
    **Validates: Requirements 1.5, 12.2**
    
    Property: Parser must extract at least one of experience or education from valid resumes
    """
    assume(len(text) > 100)  # Ensure meaningful content
    assume('EXPERIENCE' in text.upper() or 'EDUCATION' in text.upper())
    
    # Create temporary file
    temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False)
    temp_file.write(text)
    temp_file.close()
    
    try:
        # Parse sections
        identifier = SectionIdentifier()
        sections = identifier.identify_sections(text)
        
        # Verify at least one key section was identified
        key_sections = {'experience', 'education', 'skills'}
        found_sections = set(sections.keys()) & key_sections
        
        assert len(found_sections) > 0, f"No key sections found. Sections: {list(sections.keys())}"
    
    finally:
        if os.path.exists(temp_file.name):
            os.unlink(temp_file.name)


# Property 6: Section boundary identification
@pytest.mark.property
@settings(max_examples=100)
@given(text=resume_text())
def test_property_6_section_boundary_identification(text):
    """
    **Feature: talentflow-ai, Property 6: Section boundary identification**
    **Validates: Requirements 12.1**
    
    Property: Section boundaries must be correctly identified with confidence scores
    """
    assume(len(text) > 50)
    
    identifier = SectionIdentifier()
    
    # Get sections
    sections = identifier.identify_sections(text)
    
    # Get confidence scores
    confidence_scores = identifier.get_section_confidence_scores(text)
    
    # Verify sections have content
    for section_name, section_content in sections.items():
        assert isinstance(section_content, str)
        assert len(section_content) > 0
    
    # Verify confidence scores are in valid range
    for section_name, confidence in confidence_scores.items():
        assert 0.0 <= confidence <= 1.0, f"Invalid confidence {confidence} for section {section_name}"


# Property 7: Comprehensive field extraction
@pytest.mark.property
@settings(max_examples=50)
@given(text=resume_text())
def test_property_7_comprehensive_field_extraction(text):
    """
    **Feature: talentflow-ai, Property 7: Comprehensive field extraction**
    **Validates: Requirements 12.2, 12.3, 12.4**
    
    Property: All extracted entities must have confidence scores
    """
    assume(len(text) > 100)
    
    identifier = SectionIdentifier()
    extractor = EntityExtractor()
    
    sections = identifier.identify_sections(text)
    
    # Extract entities from each section
    if 'experience' in sections:
        experiences = extractor.extract_work_experience(sections['experience'])
        for exp in experiences:
            assert 'confidence' in exp
            assert 0.0 <= exp['confidence'] <= 1.0
    
    if 'education' in sections:
        educations = extractor.extract_education(sections['education'])
        for edu in educations:
            assert 'confidence' in edu
            assert 0.0 <= edu['confidence'] <= 1.0
    
    if 'skills' in sections:
        skills = extractor.extract_skills(sections['skills'])
        for skill in skills:
            assert 'confidence' in skill
            assert 0.0 <= skill['confidence'] <= 1.0


# Property 8: Format robustness
@pytest.mark.property
@settings(max_examples=50)
@given(
    text=st.text(min_size=100, max_size=1000),
    format_type=st.sampled_from(['pdf', 'docx'])
)
def test_property_8_format_robustness(text, format_type):
    """
    **Feature: talentflow-ai, Property 8: Format robustness**
    **Validates: Requirements 12.5**
    
    Property: Parser must handle various text formats without crashing
    """
    # Create temporary file
    if format_type == 'pdf':
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.pdf', delete=False)
        temp_file.close()
        
        try:
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(temp_file.name)
            c.drawString(50, 750, text[:100])
            c.save()
            
            extractor = TextExtractor()
            result = extractor.extract_text(temp_file.name)
            
            # Should either succeed or raise ValidationException
            assert isinstance(result, str) or result is None
        
        except ValidationException:
            # Expected for invalid content
            pass
        
        finally:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
    
    else:  # docx
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False)
        temp_file.close()
        
        try:
            doc = Document()
            doc.add_paragraph(text[:500])
            doc.save(temp_file.name)
            
            extractor = TextExtractor()
            result = extractor.extract_text(temp_file.name)
            
            # Should either succeed or raise ValidationException
            assert isinstance(result, str) or result is None
        
        except ValidationException:
            # Expected for invalid content
            pass
        
        finally:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)


# Property 9: Low confidence flagging
@pytest.mark.property
@settings(max_examples=50, suppress_health_check=[HealthCheck.function_scoped_fixture])
@given(file_data=docx_file())
def test_property_9_low_confidence_flagging(file_data):
    """
    **Feature: talentflow-ai, Property 9: Low confidence flagging**
    **Validates: Requirements 12.6**
    
    Property: Fields with confidence below threshold must be flagged
    """
    file_path, content = file_data
    
    try:
        parser = ResumeParser()
        parsed_resume = parser.parse_resume(file_path)
        
        # Verify low_confidence_fields is a list
        assert isinstance(parsed_resume.low_confidence_fields, list)
        
        # Verify all flagged fields reference valid paths
        for field_path in parsed_resume.low_confidence_fields:
            assert isinstance(field_path, str)
            assert len(field_path) > 0
            
            # Field path should contain valid section references
            valid_prefixes = ['work_experience', 'education', 'skills', 'certifications']
            assert any(field_path.startswith(prefix) for prefix in valid_prefixes)
    
    finally:
        if os.path.exists(file_path):
            os.unlink(file_path)
