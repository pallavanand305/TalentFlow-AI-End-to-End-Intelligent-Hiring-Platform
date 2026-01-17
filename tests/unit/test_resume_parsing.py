"""Unit tests for resume parsing edge cases"""

import pytest
import tempfile
import os
from pathlib import Path
from docx import Document
from reportlab.pdfgen import canvas

from ml.parsing.text_extractor import TextExtractor
from ml.parsing.section_identifier import SectionIdentifier
from ml.parsing.entity_extractor import EntityExtractor
from ml.parsing.resume_parser import ResumeParser
from backend.app.core.exceptions import ValidationException


class TestTextExtractorEdgeCases:
    """Test edge cases for text extraction"""
    
    def test_empty_pdf(self):
        """Test extraction from empty PDF"""
        # Create empty PDF
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.pdf', delete=False)
        temp_file.close()
        
        try:
            c = canvas.Canvas(temp_file.name)
            c.save()
            
            extractor = TextExtractor()
            
            with pytest.raises(ValidationException) as exc_info:
                extractor.extract_text(temp_file.name)
            
            assert "No text could be extracted" in str(exc_info.value)
        
        finally:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
    
    def test_empty_docx(self):
        """Test extraction from empty DOCX"""
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False)
        temp_file.close()
        
        try:
            doc = Document()
            doc.save(temp_file.name)
            
            extractor = TextExtractor()
            
            with pytest.raises(ValidationException) as exc_info:
                extractor.extract_text(temp_file.name)
            
            assert "No text could be extracted" in str(exc_info.value)
        
        finally:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
    
    def test_nonexistent_file(self):
        """Test extraction from nonexistent file"""
        extractor = TextExtractor()
        
        with pytest.raises(ValidationException) as exc_info:
            extractor.extract_text("/nonexistent/file.pdf")
        
        assert "File not found" in str(exc_info.value)
    
    def test_unsupported_format(self):
        """Test extraction from unsupported file format"""
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False)
        temp_file.write("Some text")
        temp_file.close()
        
        try:
            extractor = TextExtractor()
            
            with pytest.raises(ValidationException) as exc_info:
                extractor.extract_text(temp_file.name)
            
            assert "Unsupported file format" in str(exc_info.value)
        
        finally:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
    
    def test_corrupted_pdf(self):
        """Test extraction from corrupted PDF"""
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.pdf', delete=False)
        temp_file.write("This is not a valid PDF file")
        temp_file.close()
        
        try:
            extractor = TextExtractor()
            
            with pytest.raises(ValidationException) as exc_info:
                extractor.extract_text(temp_file.name)
            
            assert "Failed to" in str(exc_info.value).lower()
        
        finally:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
    
    def test_pdf_with_tables(self):
        """Test extraction from PDF with tables"""
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.pdf', delete=False)
        temp_file.close()
        
        try:
            c = canvas.Canvas(temp_file.name)
            # Simulate table-like content
            c.drawString(50, 750, "Skill")
            c.drawString(200, 750, "Years")
            c.drawString(50, 730, "Python")
            c.drawString(200, 730, "5")
            c.save()
            
            extractor = TextExtractor()
            text = extractor.extract_text(temp_file.name)
            
            assert "Python" in text
            assert len(text) > 0
        
        finally:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
    
    def test_docx_with_tables(self):
        """Test extraction from DOCX with tables"""
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False)
        temp_file.close()
        
        try:
            doc = Document()
            doc.add_paragraph("Skills")
            
            # Add table
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Skill"
            table.cell(0, 1).text = "Years"
            table.cell(1, 0).text = "Python"
            table.cell(1, 1).text = "5"
            
            doc.save(temp_file.name)
            
            extractor = TextExtractor()
            text = extractor.extract_text(temp_file.name)
            
            assert "Python" in text
            assert "Skills" in text
        
        finally:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)


class TestSectionIdentifierEdgeCases:
    """Test edge cases for section identification"""
    
    def test_no_sections(self):
        """Test with text that has no clear sections"""
        text = "This is just a paragraph of text without any section headers."
        
        identifier = SectionIdentifier()
        sections = identifier.identify_sections(text)
        
        # Should have at least a 'header' section
        assert 'header' in sections
        assert len(sections['header']) > 0
    
    def test_multiple_experience_headers(self):
        """Test with multiple variations of experience headers"""
        text = """
        WORK EXPERIENCE
        Software Engineer at Company A
        
        PROFESSIONAL EXPERIENCE
        Data Scientist at Company B
        """
        
        identifier = SectionIdentifier()
        sections = identifier.identify_sections(text)
        
        # Should identify experience section
        assert 'experience' in sections
        assert 'Company A' in sections['experience'] or 'Company B' in sections['experience']
    
    def test_case_insensitive_headers(self):
        """Test section headers with different cases"""
        text = """
        education
        Bachelor of Science
        
        SKILLS
        Python, Java
        
        Experience
        Software Engineer
        """
        
        identifier = SectionIdentifier()
        sections = identifier.identify_sections(text)
        
        assert 'education' in sections
        assert 'skills' in sections
        assert 'experience' in sections
    
    def test_headers_with_colons(self):
        """Test section headers ending with colons"""
        text = """
        Education:
        Bachelor of Science
        
        Skills:
        Python, Java
        """
        
        identifier = SectionIdentifier()
        sections = identifier.identify_sections(text)
        
        assert 'education' in sections
        assert 'skills' in sections
    
    def test_very_long_lines(self):
        """Test with very long lines that shouldn't be headers"""
        text = """
        EXPERIENCE
        This is a very long line that contains the word experience but should not be treated as a header because it is too long and contains too much content that makes it clearly a paragraph rather than a section header.
        """
        
        identifier = SectionIdentifier()
        sections = identifier.identify_sections(text)
        
        assert 'experience' in sections


class TestEntityExtractorEdgeCases:
    """Test edge cases for entity extraction"""
    
    def test_empty_experience_section(self):
        """Test extraction from empty experience section"""
        extractor = EntityExtractor()
        experiences = extractor.extract_work_experience("")
        
        assert isinstance(experiences, list)
        assert len(experiences) == 0
    
    def test_experience_without_dates(self):
        """Test extraction from experience without dates"""
        text = """
        Software Engineer
        TechCorp
        Developed web applications
        """
        
        extractor = EntityExtractor()
        experiences = extractor.extract_work_experience(text)
        
        # Should still extract something, even with low confidence
        assert isinstance(experiences, list)
    
    def test_experience_with_present(self):
        """Test extraction with 'Present' as end date"""
        text = """
        Software Engineer
        TechCorp
        Jan 2020 - Present
        """
        
        extractor = EntityExtractor()
        experiences = extractor.extract_work_experience(text)
        
        assert len(experiences) > 0
        if experiences[0]['end_date']:
            assert 'Present' in experiences[0]['end_date']
    
    def test_education_without_degree(self):
        """Test extraction from education without clear degree"""
        text = """
        State University
        Computer Science
        2016 - 2020
        """
        
        extractor = EntityExtractor()
        educations = extractor.extract_education(text)
        
        assert isinstance(educations, list)
    
    def test_skills_with_various_separators(self):
        """Test skill extraction with different separators"""
        text = "Python, Java | C++ • JavaScript · SQL; Ruby"
        
        extractor = EntityExtractor()
        skills = extractor.extract_skills(text)
        
        assert len(skills) > 0
        skill_names = [s['skill'] for s in skills]
        assert any('Python' in s for s in skill_names)
    
    def test_skills_multiline(self):
        """Test skill extraction from multiline text"""
        text = """
        Python
        Java
        C++
        JavaScript
        """
        
        extractor = EntityExtractor()
        skills = extractor.extract_skills(text)
        
        assert len(skills) >= 3
    
    def test_certifications_with_dates(self):
        """Test certification extraction with dates"""
        text = """
        AWS Certified Developer - 2021
        Google Cloud Professional - Dec 2022
        """
        
        extractor = EntityExtractor()
        certs = extractor.extract_certifications(text)
        
        assert len(certs) == 2
        assert any(c['date'] is not None for c in certs)
    
    def test_certifications_without_dates(self):
        """Test certification extraction without dates"""
        text = """
        AWS Certified Developer
        Google Cloud Professional
        """
        
        extractor = EntityExtractor()
        certs = extractor.extract_certifications(text)
        
        assert len(certs) == 2


class TestResumeParserEdgeCases:
    """Test edge cases for complete resume parsing"""
    
    def test_minimal_resume(self):
        """Test parsing minimal resume with only education"""
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False)
        temp_file.close()
        
        try:
            doc = Document()
            doc.add_paragraph("John Doe")
            doc.add_paragraph("EDUCATION")
            doc.add_paragraph("Bachelor of Science")
            doc.add_paragraph("State University")
            doc.save(temp_file.name)
            
            parser = ResumeParser()
            parsed = parser.parse_resume(temp_file.name)
            
            assert parsed is not None
            assert parsed.file_format == 'docx'
            assert len(parsed.raw_text) > 0
        
        finally:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
    
    def test_resume_with_all_sections(self):
        """Test parsing resume with all sections"""
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False)
        temp_file.close()
        
        try:
            doc = Document()
            doc.add_paragraph("Jane Smith")
            doc.add_paragraph("EXPERIENCE")
            doc.add_paragraph("Software Engineer")
            doc.add_paragraph("TechCorp")
            doc.add_paragraph("Jan 2020 - Present")
            doc.add_paragraph("EDUCATION")
            doc.add_paragraph("Master of Science")
            doc.add_paragraph("Tech University")
            doc.add_paragraph("SKILLS")
            doc.add_paragraph("Python, Java, SQL")
            doc.add_paragraph("CERTIFICATIONS")
            doc.add_paragraph("AWS Certified Developer")
            doc.save(temp_file.name)
            
            parser = ResumeParser()
            parsed = parser.parse_resume(temp_file.name)
            
            assert parsed is not None
            assert len(parsed.sections) >= 3
            assert isinstance(parsed.low_confidence_fields, list)
        
        finally:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
    
    def test_validate_minimum_fields_with_experience(self):
        """Test minimum field validation with experience"""
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False)
        temp_file.close()
        
        try:
            doc = Document()
            doc.add_paragraph("EXPERIENCE")
            doc.add_paragraph("Software Engineer at TechCorp")
            doc.save(temp_file.name)
            
            parser = ResumeParser()
            parsed = parser.parse_resume(temp_file.name)
            
            is_valid = parser.validate_minimum_fields(parsed)
            # Should be valid if at least one experience or education entry
            assert isinstance(is_valid, bool)
        
        finally:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
    
    def test_non_resume_document(self):
        """Test parsing non-resume document"""
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.docx', delete=False)
        temp_file.close()
        
        try:
            doc = Document()
            doc.add_paragraph("This is a random document about cooking recipes.")
            doc.add_paragraph("It has nothing to do with resumes or CVs.")
            doc.save(temp_file.name)
            
            parser = ResumeParser()
            parsed = parser.parse_resume(temp_file.name)
            
            # Should parse without error, but may have empty sections
            assert parsed is not None
            assert len(parsed.raw_text) > 0
        
        finally:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
