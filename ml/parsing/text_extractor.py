"""Text extraction from resume files"""

import os
from typing import Optional
from pathlib import Path
import pypdf
import docx

from backend.app.core.logging import get_logger
from backend.app.core.exceptions import ValidationException

logger = get_logger(__name__)


class TextExtractor:
    """Extract text from PDF and DOCX files"""
    
    SUPPORTED_EXTENSIONS = ['.pdf', '.docx']
    
    @staticmethod
    def validate_file(file_path: str) -> None:
        """
        Validate file exists and has supported extension
        
        Args:
            file_path: Path to file
        
        Raises:
            ValidationException: If file is invalid
        """
        if not os.path.exists(file_path):
            raise ValidationException(f"File not found: {file_path}")
        
        ext = Path(file_path).suffix.lower()
        if ext not in TextExtractor.SUPPORTED_EXTENSIONS:
            raise ValidationException(
                f"Unsupported file format: {ext}. Supported formats: {TextExtractor.SUPPORTED_EXTENSIONS}"
            )
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
        
        Returns:
            Extracted text
        
        Raises:
            ValidationException: If extraction fails
        """
        try:
            text_parts = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    raise ValidationException("PDF file is encrypted and cannot be processed")
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
            extracted_text = '\n'.join(text_parts)
            
            if not extracted_text.strip():
                raise ValidationException("No text could be extracted from PDF")
            
            logger.info(f"Extracted {len(extracted_text)} characters from PDF: {file_path}")
            return extracted_text
        
        except pypdf.errors.PdfReadError as e:
            logger.error(f"PDF read error: {str(e)}")
            raise ValidationException(f"Failed to read PDF file: {str(e)}")
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise ValidationException(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Extracted text
        
        Raises:
            ValidationException: If extraction fails
        """
        try:
            doc = docx.Document(file_path)
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            extracted_text = '\n'.join(text_parts)
            
            if not extracted_text.strip():
                raise ValidationException("No text could be extracted from DOCX")
            
            logger.info(f"Extracted {len(extracted_text)} characters from DOCX: {file_path}")
            return extracted_text
        
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise ValidationException(f"Failed to extract text from DOCX: {str(e)}")
    
    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """
        Extract text from resume file (PDF or DOCX)
        
        Args:
            file_path: Path to resume file
        
        Returns:
            Extracted text
        
        Raises:
            ValidationException: If file is invalid or extraction fails
        """
        # Validate file
        cls.validate_file(file_path)
        
        # Determine file type and extract
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return cls.extract_from_pdf(file_path)
        elif ext == '.docx':
            return cls.extract_from_docx(file_path)
        else:
            raise ValidationException(f"Unsupported file format: {ext}")
